"""Build PAPER3_DRAFT_v1_0.docx from the canonical source markdown.

Session 071 build pipeline. Paper 3 v1.0 source markdown is a single
file at ``docs/paper_3/source/PAPER3_DRAFT_v1_0_source.md`` covering
Abstract + Sections 1 - 10 + the Generative AI Use Declaration
appendix. The build walks the source and emits a docx mirroring its
structural choices, with the anonymized author block preserved
verbatim per the Session 069 GO 3 discipline.

Per the Paper 3 build_references.py docstring, the ACM/AISec docx
formatter was deliberately deferred to "the point at which prose
lands"; this module is that landing. Pattern is parallel to Paper 2's
``docs/paper_2/build_v1_1.py``, scoped to a single source file with a
References section appended from the verified bib entries and the
Appendix emitted last (per its own positional declaration that it
sits after the references).

Build rules:

* HTML comments (the anonymization restoration note) are stripped
  before parsing — they belong in the markdown audit trail, not the
  rendered submission docx.
* ``# Title`` → 14pt centered bold heading paragraph.
* ``**[Author block redacted for double-blind review]**`` → centered
  italic paragraph, rendered as-is.
* ``## X`` → level-1 heading.
* ``### X`` → level-2 heading.
* ``---`` decorative separators are skipped.
* Inline markdown: ``**bold**``, ``*italic*``, `` `code` `` (Consolas).
* Bullet lists (``- `` prefix) → List Bullet style.
* The Appendix section (``## Appendix:...``) is held out of source
  order and emitted AFTER the References section, matching the
  Appendix's own positional declaration.

Usage::

    python -m docs.paper_3.build_v1_0
    python -m docs.paper_3.build_v1_0 \\
        --source docs/paper_3/source/PAPER3_DRAFT_v1_0_source.md \\
        --out docs/paper_3/PAPER3_DRAFT_v1_0.docx \\
        --references docs/paper_3/references.bib
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docs.paper_3.build_references import BibEntry, parse_bib_file


# =============================================================================
# Markdown parsing helpers
# =============================================================================


_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)"
)

_HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)

_APPENDIX_HEADER_PREFIX = "## Appendix"

_AUTHOR_REDACTION_MARKER = "[Author block redacted for double-blind review]"


def strip_html_comments(text: str) -> str:
    """Remove ``<!-- ... -->`` blocks from the source markdown."""
    return _HTML_COMMENT_RE.sub("", text)


def split_appendix(source_text: str) -> tuple[str, str]:
    """Split off the Appendix section so it can be emitted after References.

    Returns ``(body, appendix)`` where ``appendix`` is the substring
    starting at the ``## Appendix`` header (inclusive). If no appendix
    is present, the second tuple element is empty.
    """
    # Find an Appendix header at the start of a line.
    pattern = re.compile(r"(?m)^## Appendix.*$")
    m = pattern.search(source_text)
    if not m:
        return source_text, ""
    return source_text[:m.start()].rstrip() + "\n", source_text[m.start():]


# =============================================================================
# Inline + block rendering
# =============================================================================


def _render_inline(paragraph, text: str) -> None:
    """Append runs to a docx paragraph honoring **bold**, *italic*, `code`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_bullet_block(non_empty_lines: list[str]) -> bool:
    return bool(non_empty_lines) and all(
        ln.lstrip().startswith("- ") for ln in non_empty_lines
    )


def _flush_block(doc: Document, block_lines: list[str]) -> None:
    """Render an accumulated block (paragraph or bullet list) into the docx."""
    non_empty = [ln for ln in block_lines if ln.strip()]
    if not non_empty:
        return
    if _is_bullet_block(non_empty):
        for ln in non_empty:
            stripped = ln.lstrip()
            if not stripped.startswith("- "):
                continue
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, stripped[2:])
        return
    p = doc.add_paragraph()
    _render_inline(p, " ".join(ln.strip() for ln in non_empty))


# =============================================================================
# Document body walker
# =============================================================================


def _emit_body_lines(doc: Document, lines: Iterable[str]) -> None:
    """Walk source markdown lines and emit headings + blocks into the doc.

    Caller is responsible for stripping HTML comments + splitting out
    the Appendix beforehand if those need separate handling.
    """
    title_emitted = False
    current_block: list[str] = []

    def flush() -> None:
        nonlocal current_block
        _flush_block(doc, current_block)
        current_block = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not title_emitted and line.startswith("# ") and not line.startswith("## "):
            flush()
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(14)
            title_emitted = True
            continue

        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.startswith("## "):
            flush()
            doc.add_heading(line[3:].strip(), level=1)
            continue

        if stripped == "---":
            flush()
            continue

        if not stripped:
            flush()
            continue

        if _AUTHOR_REDACTION_MARKER in stripped:
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _render_inline(p, stripped)
            continue

        current_block.append(line)

    flush()


# =============================================================================
# References emitter (ACM-ish single-style formatting)
# =============================================================================


def _format_reference(entry: BibEntry) -> str:
    """Render one BibEntry as a single-line reference string.

    Format is ACM-ish but compressed: ``Author. Title. Venue, Year.
    Identifier.`` Field choice is conservative; the goal is unambiguous
    attribution rather than perfect ACM Sigconf bibliography polish.
    Identifier order priority: DOI > arXiv eprint > ISBN.
    """
    f = entry.fields
    parts: list[str] = []

    author = f.get("author", "")
    if author:
        parts.append(author)

    title = f.get("title", "")
    if title:
        # Strip BibTeX brace-protection markers for plain rendering.
        title_plain = title.replace("{", "").replace("}", "")
        parts.append(title_plain + ".")

    venue = f.get("journal") or f.get("booktitle") or ""
    if venue:
        parts.append(venue + ".")

    publisher = f.get("publisher", "")
    if publisher and publisher not in venue:
        parts.append(publisher + ".")

    pages = f.get("pages", "")
    if pages:
        parts.append(f"pp. {pages}.")

    year = f.get("year", "")
    if year:
        parts.append(f"({year}).")

    if f.get("doi"):
        parts.append(f"DOI: {f['doi']}.")
    elif f.get("eprint"):
        archive = f.get("archiveprefix", "arXiv")
        parts.append(f"{archive}:{f['eprint']}.")
    elif f.get("isbn"):
        parts.append(f"ISBN: {f['isbn']}.")

    return " ".join(parts)


def _add_references_section(doc: Document, entries: tuple[BibEntry, ...]) -> None:
    """Append a References section listing every entry.

    Order is by bibkey (deterministic). Each entry is rendered as one
    paragraph with the bibkey in bold prefix brackets for traceability
    with the prose `(Author, Year)` form.
    """
    if not entries:
        return
    doc.add_heading("References", level=1)
    for entry in sorted(entries, key=lambda e: e.key):
        p = doc.add_paragraph()
        prefix_run = p.add_run(f"[{entry.key}] ")
        prefix_run.bold = True
        p.add_run(_format_reference(entry))


# =============================================================================
# Document assembly
# =============================================================================


def build_document(source_text: str, bib_path: Path) -> Document:
    """Build the Paper 3 v1.0 document.

    Steps: strip HTML comments → split off appendix → walk body →
    append References → append Appendix.
    """
    cleaned = strip_html_comments(source_text)
    body, appendix = split_appendix(cleaned)

    doc = Document()
    _emit_body_lines(doc, body.splitlines())

    if bib_path.exists():
        entries = parse_bib_file(bib_path)
        _add_references_section(doc, entries)

    if appendix.strip():
        _emit_body_lines(doc, appendix.splitlines())

    return doc


# =============================================================================
# CLI
# =============================================================================


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Build Paper 3 v1.0 docx from canonical source markdown.",
    )
    parser.add_argument(
        "--source", type=Path,
        default=Path("docs/paper_3/source/PAPER3_DRAFT_v1_0_source.md"),
    )
    parser.add_argument(
        "--out", type=Path,
        default=Path("docs/paper_3/PAPER3_DRAFT_v1_0.docx"),
    )
    parser.add_argument(
        "--references", type=Path,
        default=Path("docs/paper_3/references.bib"),
    )
    return parser


def main(argv: Optional[Iterable[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(None if argv is None else list(argv))

    source_text = args.source.resolve().read_text(encoding="utf-8")
    doc = build_document(source_text, args.references.resolve())
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"[V1.0 BUILD] wrote {args.out}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
