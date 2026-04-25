"""Build PAPER2_DRAFT_v1_1.docx by integrating the v1.1 prose source
into the section/figure structure established by ``build_skeleton.py``.

The v1 skeleton (``PAPER2_DRAFT_v1.docx``) is the historical artifact
from Session 051 and is not modified. The v1.1 build mirrors the
skeleton's section ordering, heading levels, and figure placements,
but substitutes the [TODO] markers with prose pulled from the v1.1
markdown source.

Usage::

    python -m docs.paper_2.build_v1_1
    python -m docs.paper_2.build_v1_1 \\
        --source docs/paper_2/source/PAPER2_DRAFT_v1_1_source.md \\
        --out docs/paper_2/PAPER2_DRAFT_v1_1.docx \\
        --figures-dir docs/paper_2/figures
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from docs.paper_2.build_skeleton import (
    EXPECTED_SUBSECTION_HEADERS,
    EXPECTED_TOP_LEVEL_HEADERS,
    FIGURE_CAPTIONS,
)


# =============================================================================
# Markdown source parser
# =============================================================================


@dataclass(frozen=True)
class ParsedSource:
    """Top matter + section_id -> list of block-level markdown strings."""

    title: str
    author: str
    affiliation: str
    date: str
    sections: dict[str, list[str]]


# Section IDs are normalized to a canonical key:
#   "## Abstract"                      -> "abstract"
#   "## 1. Introduction"               -> "1"
#   "### 5.1 Design"                   -> "5.1"
#   "### 6.2 Findings 9 and 10"        -> "6.2"
#   "## References"                    -> "references"
_NUMBERED_HEADER_RE = re.compile(r"^(\d+(?:\.\d+)?)\.?\s+(.*)$")


def _normalize_section_id(header_text: str) -> str:
    """Pull the canonical section id out of a header string."""
    stripped = header_text.strip()
    m = _NUMBERED_HEADER_RE.match(stripped)
    if m:
        return m.group(1)
    return stripped.lower().replace(" ", "_")


def parse_source(md_path: Path) -> ParsedSource:
    """Parse the v1.1 source markdown into a ParsedSource."""
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()

    # Top matter: title is the first '# ' line; author/affiliation/date
    # follow as plain lines until the first '---'.
    title = ""
    author = ""
    affiliation = ""
    date = ""

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if line.startswith("# "):
            title = line[2:].strip()
            idx += 1
            break
        idx += 1

    # Collect non-empty lines until we hit '---' or '##'
    top_matter: list[str] = []
    while idx < len(lines):
        line = lines[idx].rstrip()
        if line.startswith("##") or line.startswith("---"):
            break
        if line.strip():
            top_matter.append(line.strip())
        idx += 1
    # Conventional ordering: author, affiliation, date.
    if len(top_matter) >= 1:
        author = top_matter[0].replace("**", "").strip()
    if len(top_matter) >= 2:
        affiliation = top_matter[1].strip()
    if len(top_matter) >= 3:
        date = top_matter[2].strip()

    # Section parsing.
    sections: dict[str, list[str]] = {}
    current_id: Optional[str] = None
    current_blocks: list[str] = []
    current_block_lines: list[str] = []

    def _flush_block() -> None:
        if current_block_lines:
            block = "\n".join(current_block_lines).strip()
            if block:
                current_blocks.append(block)
            current_block_lines.clear()

    def _flush_section() -> None:
        nonlocal current_id, current_blocks
        if current_id is not None:
            sections.setdefault(current_id, []).extend(current_blocks)
        current_blocks = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Section divider — terminates current block, ignored otherwise.
        if stripped == "---":
            _flush_block()
            idx += 1
            continue

        # Headings establish or change the current section/subsection.
        if stripped.startswith("### "):
            _flush_block()
            _flush_section()
            current_id = _normalize_section_id(stripped[4:])
            current_blocks = []
            idx += 1
            continue
        if stripped.startswith("## "):
            _flush_block()
            _flush_section()
            current_id = _normalize_section_id(stripped[3:])
            current_blocks = []
            idx += 1
            continue

        # Blank line ends the current block.
        if stripped == "":
            _flush_block()
            idx += 1
            continue

        current_block_lines.append(line.rstrip())
        idx += 1

    _flush_block()
    _flush_section()

    return ParsedSource(
        title=title,
        author=author,
        affiliation=affiliation,
        date=date,
        sections=sections,
    )


# =============================================================================
# Inline markdown -> docx runs
# =============================================================================


# Tokens we care about, in priority order. Bold (**) must be tested
# before italic (*) so the longer marker wins.
_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)"
)


def _render_inline(paragraph, text: str) -> None:
    """Append runs to a docx paragraph, honoring **bold**, *italic*, `code`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_prose_block(doc: Document, block_md: str) -> None:
    """Render one markdown block (paragraph or bullet list) into the docx."""
    lines = block_md.splitlines()
    is_bullet_block = lines and all(
        ln.lstrip().startswith("- ") for ln in lines if ln.strip()
    )
    if is_bullet_block:
        for ln in lines:
            stripped = ln.lstrip()
            if not stripped.startswith("- "):
                continue
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, stripped[2:])
        return

    # Single-paragraph block: join lines with a space (markdown semantics).
    p = doc.add_paragraph()
    _render_inline(p, " ".join(ln.strip() for ln in lines if ln.strip()))


def _add_section_prose(
    doc: Document,
    sections: dict[str, list[str]],
    section_id: str,
) -> None:
    """Emit every block stored under ``section_id``. Empty sections are
    silently skipped (sections 5, 6, 7 in the v1.1 source are headers
    only, with content living in their .1 / .2 / .3 subsections)."""
    blocks = sections.get(section_id)
    if not blocks:
        return
    for block in blocks:
        _add_prose_block(doc, block)


# =============================================================================
# Docx primitives (mirroring build_skeleton.py conventions)
# =============================================================================


def _add_header(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_figure_placeholder(
    doc: Document,
    figure_path: Path,
    caption: str,
    max_width_inches: float = 5.8,
) -> None:
    doc.add_picture(str(figure_path), width=Inches(max_width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)


# =============================================================================
# Document assembly
# =============================================================================


def build_document(
    source: ParsedSource,
    figures_dir: Path,
) -> Document:
    """Build the v1.1 document, mirroring the v1 skeleton structure."""
    doc = Document()

    # Title + author block (pulled from source, NOT from skeleton).
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(source.title)
    run.bold = True
    run.font.size = Pt(14)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_block_lines = [source.author, source.affiliation, source.date]
    author.add_run("\n".join(line for line in author_block_lines if line))

    keywords_p = doc.add_paragraph()
    keywords_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kw_run = keywords_p.add_run(
        "Keywords: Large Language Models; Prompt Injection; "
        "Adversarial Robustness; Deterministic Reasoning; Cybersecurity"
    )
    kw_run.italic = True
    kw_run.font.size = Pt(10)

    doc.add_paragraph()

    # Abstract.
    _add_header(doc, "Abstract", level=1)
    _add_section_prose(doc, source.sections, "abstract")

    # 1. Introduction.
    _add_header(doc, "1. Introduction", level=1)
    _add_section_prose(doc, source.sections, "1")

    # 2. Related Work.
    _add_header(doc, "2. Related Work", level=1)
    _add_section_prose(doc, source.sections, "2")

    # 3. Background: ARES Architecture (figure 1 inlined here).
    _add_header(doc, "3. Background: ARES Architecture", level=1)
    _add_section_prose(doc, source.sections, "3")
    _add_figure_placeholder(
        doc, figures_dir / "fig1_architecture.png", FIGURE_CAPTIONS["fig1"],
    )

    # 4. Threat Model & Corpus. Source markdown splits this into 4.1 and
    # 4.2 subsections; the v1 skeleton's TODO body collapsed both. We
    # honor the source structure and emit the subsections explicitly.
    _add_header(doc, "4. Threat Model & Corpus", level=1)
    _add_section_prose(doc, source.sections, "4")
    _add_header(doc, "4.1 Threat Model", level=2)
    _add_section_prose(doc, source.sections, "4.1")
    _add_header(doc, "4.2 Corpus Construction", level=2)
    _add_section_prose(doc, source.sections, "4.2")

    # 5. Syntactic Firewall (5.1, 5.2 with figure 2 after 5.2).
    _add_header(doc, "5. Syntactic Firewall", level=1)
    _add_section_prose(doc, source.sections, "5")
    _add_header(doc, "5.1 Design", level=2)
    _add_section_prose(doc, source.sections, "5.1")
    _add_header(doc, "5.2 Findings 7 & 8", level=2)
    _add_section_prose(doc, source.sections, "5.2")
    _add_figure_placeholder(
        doc, figures_dir / "fig2_firewall_detection.png",
        FIGURE_CAPTIONS["fig2"],
    )

    # 6. Skeptic Ablation (6.1 + 6.2).
    _add_header(doc, "6. Skeptic Ablation", level=1)
    _add_section_prose(doc, source.sections, "6")
    _add_header(doc, "6.1 Methodology & Pre-Registered Rubric", level=2)
    _add_section_prose(doc, source.sections, "6.1")
    _add_header(doc, "6.2 Findings 9 & 10", level=2)
    _add_section_prose(doc, source.sections, "6.2")

    # 7. The Light Skeptic (7.1, 7.2, 7.3 with figures 3, 4, 5 after 7.3).
    _add_header(doc, "7. The Light Skeptic", level=1)
    _add_section_prose(doc, source.sections, "7")
    _add_header(doc, "7.1 Four-Rule Engine", level=2)
    _add_section_prose(doc, source.sections, "7.1")
    _add_header(doc, "7.2 Three-Way Benchmark", level=2)
    _add_section_prose(doc, source.sections, "7.2")
    _add_header(doc, "7.3 Finding 11", level=2)
    _add_section_prose(doc, source.sections, "7.3")
    _add_figure_placeholder(
        doc, figures_dir / "fig3_family_heatmap.png", FIGURE_CAPTIONS["fig3"],
    )
    _add_figure_placeholder(
        doc, figures_dir / "fig4_scenario_verdicts.png",
        FIGURE_CAPTIONS["fig4"],
    )
    _add_figure_placeholder(
        doc, figures_dir / "fig5_rubric_bands.png", FIGURE_CAPTIONS["fig5"],
    )

    # 8 - 11.
    _add_header(doc, "8. Discussion", level=1)
    _add_section_prose(doc, source.sections, "8")

    _add_header(doc, "9. Limitations", level=1)
    _add_section_prose(doc, source.sections, "9")

    _add_header(doc, "10. Future Work", level=1)
    _add_section_prose(doc, source.sections, "10")

    _add_header(doc, "11. Conclusion", level=1)
    _add_section_prose(doc, source.sections, "11")

    # References — header only; build_references.py fills the entries.
    _add_header(doc, "References", level=1)

    return doc


# =============================================================================
# Section IDs the v1.1 build relies on (for test assertions)
# =============================================================================


REQUIRED_SECTION_IDS: tuple[str, ...] = (
    "abstract",
    "1", "2", "3",
    "4", "4.1", "4.2",
    "5", "5.1", "5.2",
    "6", "6.1", "6.2",
    "7", "7.1", "7.2", "7.3",
    "8", "9", "10", "11",
)


# =============================================================================
# CLI
# =============================================================================


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Integrate Paper 2 v1.1 prose into the v1 skeleton.",
    )
    parser.add_argument(
        "--source", type=Path,
        default=Path("docs/paper_2/source/PAPER2_DRAFT_v1_1_source.md"),
    )
    parser.add_argument(
        "--out", type=Path,
        default=Path("docs/paper_2/PAPER2_DRAFT_v1_1.docx"),
    )
    parser.add_argument(
        "--figures-dir", type=Path,
        default=Path("docs/paper_2/figures"),
    )
    return parser


def main(argv: Optional[Iterable[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(None if argv is None else list(argv))

    source = parse_source(args.source.resolve())
    doc = build_document(source, args.figures_dir.resolve())
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"[V1.1 BUILD] wrote {args.out}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
