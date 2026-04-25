"""Compile docs/paper_2/references.bib into the References section
of docs/paper_2/PAPER2_DRAFT_v1_1.docx (ACM/AISec author-year format).

The script appends formatted reference entries to the existing
``References`` heading produced by ``build_v1_1.py``. Citation keys
in the .bib file must match the in-text citations in the v1.1 prose
source.

Usage::

    python -m docs.paper_2.build_references
    python -m docs.paper_2.build_references \\
        --bib docs/paper_2/references.bib \\
        --docx docs/paper_2/PAPER2_DRAFT_v1_1.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.shared import Pt


# =============================================================================
# BibTeX parser (minimal — handles the entry shapes used by references.bib)
# =============================================================================


@dataclass(frozen=True)
class BibEntry:
    """One bibliography entry."""

    entry_type: str
    key: str
    fields: dict[str, str]

    def get(self, name: str, default: str = "") -> str:
        return self.fields.get(name, default)


_ENTRY_HEADER_RE = re.compile(
    r"@(\w+)\s*\{\s*([^,\s]+)\s*,",
    re.IGNORECASE,
)


def _scan_balanced(text: str, start: int, opener: str = "{", closer: str = "}") -> int:
    """Return the index AFTER the matching closer for the opener at start."""
    if text[start] != opener:
        raise ValueError(
            f"_scan_balanced: expected '{opener}' at index {start}, "
            f"got '{text[start]!r}'"
        )
    depth = 1
    i = start + 1
    while i < len(text) and depth > 0:
        c = text[i]
        if c == opener:
            depth += 1
        elif c == closer:
            depth -= 1
        i += 1
    if depth != 0:
        raise ValueError(
            f"_scan_balanced: unbalanced braces starting at {start}"
        )
    return i  # index AFTER the matching closer


def _parse_fields(body: str) -> dict[str, str]:
    """Parse the inside of an entry into field -> value."""
    fields: dict[str, str] = {}
    i = 0
    n = len(body)
    while i < n:
        # Skip whitespace and commas.
        while i < n and body[i] in " \t\r\n,":
            i += 1
        if i >= n:
            break
        # Read field name.
        m = re.match(r"(\w+)\s*=\s*", body[i:])
        if not m:
            # Not a field line; skip to next comma/newline.
            while i < n and body[i] != "\n":
                i += 1
            continue
        name = m.group(1).lower()
        i += m.end()
        if i >= n:
            break
        # Read value: braced {...} or quoted "..." or bare.
        if body[i] == "{":
            close = _scan_balanced(body, i, "{", "}")
            value = body[i + 1:close - 1]
            i = close
        elif body[i] == '"':
            close = body.find('"', i + 1)
            if close < 0:
                break
            value = body[i + 1:close]
            i = close + 1
        else:
            # Bare value — read until comma or newline.
            j = i
            while j < n and body[j] not in ",\n":
                j += 1
            value = body[i:j].strip()
            i = j
        fields[name] = " ".join(value.split())
    return fields


def parse_bib(text: str) -> tuple[BibEntry, ...]:
    """Parse a BibTeX-format string into a tuple of BibEntry."""
    entries: list[BibEntry] = []
    pos = 0
    while True:
        m = _ENTRY_HEADER_RE.search(text, pos)
        if not m:
            break
        entry_type = m.group(1).lower()
        key = m.group(2).strip()
        # Find the opening brace of the entry body.
        body_open = text.find("{", m.start())
        body_close = _scan_balanced(text, body_open, "{", "}")
        # The header itself opened at body_open; everything inside is fields.
        # m.end() lands right after the first comma; that's where fields begin.
        fields_text = text[m.end():body_close - 1]
        fields = _parse_fields(fields_text)
        entries.append(BibEntry(entry_type=entry_type, key=key, fields=fields))
        pos = body_close
    return tuple(entries)


def parse_bib_file(path: Path) -> tuple[BibEntry, ...]:
    return parse_bib(path.read_text(encoding="utf-8"))


# =============================================================================
# ACM/AISec author-year formatting
# =============================================================================


_LATEX_ACCENT_RE = re.compile(r"\{\\['`^\"~=.]([a-zA-Z])\}")
_LATEX_CASEPROT_RE = re.compile(r"\{([^{}\\]+)\}")


def _strip_latex(text: str) -> str:
    """Cheap LaTeX renderer. Handles two BibTeX conventions:

      * Accent escapes:           ``Fr{\\'e}d{\\'e}ric`` -> ``Frederic``
      * Bare case-protect braces: ``{OWASP} Top 10``     -> ``OWASP Top 10``
    """
    out = _LATEX_ACCENT_RE.sub(r"\1", text)
    # Run repeatedly so nested-but-disjoint braces are collapsed.
    while _LATEX_CASEPROT_RE.search(out):
        out = _LATEX_CASEPROT_RE.sub(r"\1", out)
    return out


def _format_authors(raw: str) -> str:
    """Convert BibTeX 'A and B and C' into 'A, B, and C'."""
    if not raw:
        return ""
    parts = [p.strip() for p in re.split(r"\s+and\s+", _strip_latex(raw))]
    parts = [p for p in parts if p]
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]} and {parts[1]}"
    return ", ".join(parts[:-1]) + f", and {parts[-1]}"


def format_entry(entry: BibEntry) -> str:
    """Render one BibEntry in ACM/AISec author-year format."""
    authors = _format_authors(entry.get("author"))
    year = entry.get("year")
    title = _strip_latex(entry.get("title"))
    venue = _strip_latex(
        entry.get("journal") or entry.get("booktitle") or entry.get("institution")
        or ""
    )

    pieces: list[str] = []
    if authors:
        pieces.append(authors + ".")
    if year:
        pieces.append(year + ".")
    if title:
        pieces.append(title + ".")
    if venue:
        pieces.append(venue + ".")
    if entry.get("eprint"):
        prefix = entry.get("archivePrefix", "arXiv")
        pieces.append(f"{prefix}:{entry.get('eprint')}.")
    if entry.get("doi"):
        pieces.append(f"DOI:{entry.get('doi')}.")
    if entry.get("note"):
        pieces.append(entry.get("note"))
    return " ".join(pieces).strip()


# =============================================================================
# Docx integration
# =============================================================================


def _find_references_paragraph_index(doc: Document) -> int:
    """Locate the 'References' heading in the doc. Returns body index."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip() == "References":
            return i
    raise LookupError("References heading not found in document")


def append_references(
    doc: Document,
    entries: Iterable[BibEntry],
    sort: bool = True,
) -> int:
    """Append formatted reference entries below the References heading.

    Returns the number of entries written.
    """
    # Validate the heading exists (raises if not).
    _find_references_paragraph_index(doc)

    ordered = sorted(entries, key=lambda e: e.key) if sort else list(entries)
    for entry in ordered:
        text = format_entry(entry)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        # Track entry key as a marker for downstream tooling / tests.
        marker = p.add_run(f" [{entry.key}]")
        marker.italic = True
        marker.font.size = Pt(8)
    return len(ordered)


def compile_references(
    docx_path: Path,
    bib_path: Path,
    out_path: Optional[Path] = None,
) -> Path:
    """Read the docx, append references, write back (or to out_path)."""
    if not docx_path.exists():
        raise FileNotFoundError(
            f"Target docx not found: {docx_path}. Run build_v1_1 first."
        )
    if not bib_path.exists():
        raise FileNotFoundError(f"Bibliography not found: {bib_path}.")
    doc = Document(str(docx_path))
    entries = parse_bib_file(bib_path)
    if not entries:
        raise ValueError(f"No bibliography entries parsed from {bib_path}")
    append_references(doc, entries)
    target = out_path or docx_path
    doc.save(str(target))
    return target


# =============================================================================
# In-text citation resolution (used by tests)
# =============================================================================


# Match ``(Author[, ...] (et al.)?, YYYY)`` with one optional pre-author
# qualifier word (Berdoz, Rugli, and Wattenhofer style).
_CITATION_RE = re.compile(
    r"\(([A-Z][A-Za-z\-]+(?:,\s+[A-Za-z\-]+(?:,\s+and\s+[A-Za-z\-]+)?)?"
    r"(?:\s+et\s+al\.)?)\s*,\s*(\d{4})\)"
)


def extract_citations(prose: str) -> tuple[tuple[str, str], ...]:
    """Return the unique ``(author_token, year)`` citations found in prose."""
    seen: set[tuple[str, str]] = set()
    out: list[tuple[str, str]] = []
    for m in _CITATION_RE.finditer(prose):
        key = (m.group(1).strip(), m.group(2))
        if key not in seen:
            seen.add(key)
            out.append(key)
    return tuple(out)


def citation_to_bibkey(citation: tuple[str, str]) -> str:
    """Map a (author_token, year) citation to a canonical bib key.

    Examples:
        ("Gmys-Casiano", "2026") -> "gmys-casiano-2026"
        ("Hossain et al.", "2025") -> "hossain-2025"
        ("Berdoz, Rugli, and Wattenhofer", "2026") ->
            "berdoz-rugli-wattenhofer-2026"
        ("OWASP", "2025") -> "owasp-2025"
    """
    author, year = citation
    cleaned = author.replace(" et al.", "")
    parts = [p.strip().lower().replace(" ", "-") for p in cleaned.split(",")]
    parts = [re.sub(r"^and-", "", p) for p in parts if p]
    return "-".join(parts + [year])


# =============================================================================
# CLI
# =============================================================================


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Compile references.bib into the v1.1 docx.",
    )
    parser.add_argument(
        "--bib", type=Path,
        default=Path("docs/paper_2/references.bib"),
    )
    parser.add_argument(
        "--docx", type=Path,
        default=Path("docs/paper_2/PAPER2_DRAFT_v1_1.docx"),
    )
    parser.add_argument(
        "--out", type=Path, default=None,
        help="Write to this path instead of overwriting --docx.",
    )
    return parser


def main(argv: Optional[Iterable[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(None if argv is None else list(argv))

    target = compile_references(
        docx_path=args.docx.resolve(),
        bib_path=args.bib.resolve(),
        out_path=args.out.resolve() if args.out else None,
    )
    print(f"[REFS] wrote {target}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
