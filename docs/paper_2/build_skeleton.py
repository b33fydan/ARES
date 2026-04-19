"""Builds the Paper 2 docx skeleton with section headers, figure
placeholders, and caption stubs. Zero prose — every section body is a
TODO marker. Run this script to regenerate PAPER2_DRAFT_v1.docx.

Usage:
    python -m docs.paper_2.build_skeleton
    python -m docs.paper_2.build_skeleton \
        --out docs/paper_2/PAPER2_DRAFT_v1.docx \
        --figures-dir docs/paper_2/figures
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


TITLE = (
    "The Deterministic Skeptic: Four Rules Match an LLM Agent in "
    "Adversarial Cybersecurity Threat Analysis"
)
AUTHOR = "Daniel Corrigan"
AFFILIATION = "Skyframe Innovations"
KEYWORDS = (
    "Large Language Models; Prompt Injection; Adversarial Robustness; "
    "Deterministic Reasoning; Cybersecurity"
)


# Caption stubs carry every numerical claim we will defend with
# number_check.py. Keep them in sync with default_claims() there.
FIGURE_CAPTIONS: dict[str, str] = {
    "fig1": (
        "Figure 1. ARES pipeline — three Skeptic variants share the same "
        "frame. Architect (LLM) and OracleJudge (deterministic) are "
        "unchanged across Sessions 048, 049, and 050; only the Skeptic "
        "step swaps between LLM, null (ablation), and Light (four rules)."
    ),
    "fig2": (
        "Figure 2. Session 048 live benchmark (claude-sonnet-4-6, n=27): "
        "direct detection 1.00 (n=4), framing detection 0.00 (n=19), "
        "propagation detection 0.75 (n=4); verdict accuracies 0.75 / 0.79 "
        "/ 0.75."
    ),
    "fig3": (
        "Figure 3. Per-family verdict accuracy across three pipeline "
        "variants (Session 050, n=15 strategies across 5 families). "
        "Light matches full on severity (1.00), temporal (1.00), "
        "causal (1.00), and narrative (0.75); ties on authority (0.833)."
    ),
    "fig4": (
        "Figure 4. Session 050 per-scenario verdicts across the three "
        "pipeline variants (n=25 framing scenarios). Light agrees with "
        "full on 21/25 scenarios; light rescues INJ-025 where full "
        "over-moderated, and loses INJ-008 where full correctly confirmed."
    ),
    "fig5": (
        "Figure 5. Pre-registered Finding-11 rubric bands. Landed values: "
        "full = 0.8400, light = 0.8400, ablated = 0.7200 on n=25. Delta "
        "(light − full) = 0.0000. Verdict: SUPPORTED."
    ),
}


def _add_header(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_todo(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[TODO — {text}]")
    run.italic = True


def _add_figure_placeholder(
    doc: Document,
    figure_path: Path,
    caption: str,
    max_width_inches: float = 5.8,
) -> None:
    """Inline the PNG and then the caption. The caption appears in italics."""
    doc.add_picture(str(figure_path), width=Inches(max_width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)


def build_document(figures_dir: Path) -> Document:
    doc = Document()

    # Title + author block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run(f"{AUTHOR}\n{AFFILIATION}")

    keywords_p = doc.add_paragraph()
    keywords_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kw_run = keywords_p.add_run(f"Keywords: {KEYWORDS}")
    kw_run.italic = True
    kw_run.font.size = Pt(10)

    doc.add_paragraph()

    # Abstract
    _add_header(doc, "Abstract", level=1)
    _add_todo(doc, "200-word abstract. Cover: Finding 7 (firewall blind to "
                    "framing), Finding 9 (Skeptic adds ~10pp), Finding 10 "
                    "(OracleJudge dismiss gate is Skeptic-gated), "
                    "Finding 11 (deterministic 4-rule Light Skeptic matches "
                    "the LLM Skeptic at 0.84 framing accuracy on n=25).")

    # 1. Introduction
    _add_header(doc, "1. Introduction", level=1)
    _add_todo(doc, "Introduce the claim: LLM Skeptic's contribution to "
                    "verdict accuracy is replaceable by four evidence-graph "
                    "rules at zero inference cost. Cite Paper 1 for the "
                    "single-turn pivot and the ARES architecture.")

    # 2. Related Work
    _add_header(doc, "2. Related Work", level=1)
    _add_todo(doc, "Dialectical LLM architectures; prompt injection "
                    "taxonomies; deterministic firewalls; rule-based vs "
                    "neural Skepticism.")

    # 3. Background: ARES Architecture
    _add_header(doc, "3. Background: ARES Architecture", level=1)
    _add_todo(doc, "Summarize Architect / Firewall / Skeptic / OracleJudge. "
                    "Reference Paper 1 for the full derivation.")
    _add_figure_placeholder(
        doc, figures_dir / "fig1_architecture.png", FIGURE_CAPTIONS["fig1"],
    )

    # 4. Threat Model & Corpus
    _add_header(doc, "4. Threat Model & Corpus", level=1)
    _add_todo(doc, "Describe the 33-scenario injection corpus: 4 direct, "
                    "25 framing (22 seed+expansion + 3 temporal expansion), "
                    "4 propagation. Category B framing families: severity, "
                    "authority, temporal, causal, narrative.")

    # 5. Syntactic Firewall
    _add_header(doc, "5. Syntactic Firewall", level=1)
    _add_header(doc, "5.1 Design", level=2)
    _add_todo(doc, "Compile regex set + evidence-graph authorization "
                    "check. Explain why it runs first in the pipeline.")
    _add_header(doc, "5.2 Findings 7 & 8", level=2)
    _add_todo(doc, "Finding 7: deterministic firewall catches 100% of "
                    "direct injections, 0% of framing. Finding 8: sanitizing "
                    "tainted output isn't sufficient against evidence-"
                    "embedded framing.")
    _add_figure_placeholder(
        doc, figures_dir / "fig2_firewall_detection.png",
        FIGURE_CAPTIONS["fig2"],
    )

    # 6. Skeptic Ablation
    _add_header(doc, "6. Skeptic Ablation", level=1)
    _add_header(doc, "6.1 Methodology & Pre-Registered Rubric", level=2)
    _add_todo(doc, "Session 049 protocol. Pre-registered Finding-9 rubric: "
                    "ablated<0.55 → SUPPORTED; 0.55≤ablated<0.70 → "
                    "AMBIGUOUS; ablated≥0.70 → NOT SUPPORTED.")
    _add_header(doc, "6.2 Findings 9 & 10", level=2)
    _add_todo(doc, "Finding 9: ablated accuracy 0.7200 (n=25) landed in "
                    "AMBIGUOUS band. Finding 10: OracleJudge decision "
                    "table structurally requires skeptic_confidence≥0.70 "
                    "for THREAT_DISMISSED — ablation cannot reach that "
                    "verdict.")

    # 7. The Light Skeptic
    _add_header(doc, "7. The Light Skeptic", level=1)
    _add_header(doc, "7.1 Four-Rule Engine", level=2)
    _add_todo(doc, "R1 authorization_marker_present (weight 0.4); "
                    "R2 benign_explanation_marker_present (weight 0.3); "
                    "R3 kill_chain_stage_low (weight 0.2 + malign cap 0.5); "
                    "R4 default_floor (confidence 0.5). Final confidence = "
                    "clamp(0.5 + benign - malign, 0, 1).")
    _add_header(doc, "7.2 Three-Way Benchmark", level=2)
    _add_todo(doc, "Session 050 setup: 25 framing scenarios, three "
                    "variants, Session 048/049 reuse to conserve LLM cost. "
                    "Session 050 runner: run_three_way_benchmark.py.")
    _add_header(doc, "7.3 Finding 11", level=2)
    _add_todo(doc, "SUPPORTED. Full 0.8400 vs Light 0.8400 (Δ = 0.0000). "
                    "Ablated 0.7200. Per-family match on severity / "
                    "temporal / causal / narrative; authority ties at "
                    "0.833. Two disagreements (INJ-008 lost, INJ-025 "
                    "rescued) cancel out at aggregate.")
    _add_figure_placeholder(
        doc, figures_dir / "fig3_family_heatmap.png", FIGURE_CAPTIONS["fig3"],
    )
    _add_figure_placeholder(
        doc, figures_dir / "fig4_scenario_verdicts.png",
        FIGURE_CAPTIONS["fig4"],
    )
    _add_figure_placeholder(
        doc, figures_dir / "fig5_rubric_bands.png", FIGURE_CAPTIONS["fig5"],
    )

    # 8. Discussion
    _add_header(doc, "8. Discussion", level=1)
    _add_todo(doc, "Cost implications: Light eliminates one LLM round-trip "
                    "per cycle. Explainability: each rule produces a "
                    "rationale entry. Publishability: Finding 11 is a "
                    "replaceability claim, not a dominance claim.")

    # 9. Limitations
    _add_header(doc, "9. Limitations", level=1)
    _add_todo(doc, "Single model (claude-sonnet-4-6); single language "
                    "(English); 33 synthetic scenarios authored by one "
                    "researcher. Rule set is v1 — malign-side rules not "
                    "yet populated.")

    # 10. Future Work
    _add_header(doc, "10. Future Work", level=1)
    _add_todo(doc, "Malign-side rules (Session 051 recommendation #2). "
                    "Per-model scaling (Opus 4.7, Haiku 4.5). Multi-lingual "
                    "framing corpus. Production deployment telemetry.")

    # 11. Conclusion
    _add_header(doc, "11. Conclusion", level=1)
    _add_todo(doc, "Restate Finding 11 in one sentence; tie back to the "
                    "ARES architecture and Paper 1's single-turn pivot.")

    # References
    _add_header(doc, "References", level=1)
    _add_todo(doc, "Paper 1 self-citation. Related dialectical LLM "
                    "papers. MITRE ATT&CK references for each threat "
                    "category in the corpus.")

    return doc


# Section headers (top-level + subsection) for test assertions.
EXPECTED_TOP_LEVEL_HEADERS: tuple[str, ...] = (
    "Abstract",
    "1. Introduction",
    "2. Related Work",
    "3. Background: ARES Architecture",
    "4. Threat Model & Corpus",
    "5. Syntactic Firewall",
    "6. Skeptic Ablation",
    "7. The Light Skeptic",
    "8. Discussion",
    "9. Limitations",
    "10. Future Work",
    "11. Conclusion",
    "References",
)

EXPECTED_SUBSECTION_HEADERS: tuple[str, ...] = (
    "5.1 Design",
    "5.2 Findings 7 & 8",
    "6.1 Methodology & Pre-Registered Rubric",
    "6.2 Findings 9 & 10",
    "7.1 Four-Rule Engine",
    "7.2 Three-Way Benchmark",
    "7.3 Finding 11",
)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Build the Paper 2 docx skeleton",
    )
    parser.add_argument(
        "--out", type=Path,
        default=Path("docs/paper_2/PAPER2_DRAFT_v1.docx"),
    )
    parser.add_argument(
        "--figures-dir", type=Path,
        default=Path("docs/paper_2/figures"),
    )
    return parser


def main(argv: Optional[Iterable[str]] = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(None if argv is None else list(argv))

    doc = build_document(args.figures_dir.resolve())
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"[SKELETON] wrote {args.out}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
