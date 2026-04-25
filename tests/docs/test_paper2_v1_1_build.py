"""Tests for docs/paper_2/build_v1_1.py.

Covers:
    * Markdown source parsing produces expected section keys.
    * Inline run rendering honors **bold**, *italic*, `code`.
    * End-to-end build on a fixture skeleton + fixture markdown.
    * End-to-end build on the real v1.1 source produces a populated docx.
    * Title / author / heading / figure invariants.
"""

from __future__ import annotations

from pathlib import Path

import pytest


REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = REPO_ROOT / "docs/paper_2/source/PAPER2_DRAFT_v1_1_source.md"
FIGURES_DIR = REPO_ROOT / "docs/paper_2/figures"


# ----------------------------------------------------------------------------
# Markdown source parser
# ----------------------------------------------------------------------------


class TestParseSource:
    def test_parses_real_source(self):
        from docs.paper_2.build_v1_1 import parse_source
        parsed = parse_source(SOURCE_MD)
        assert parsed.title
        assert "Deterministic Skeptic" in parsed.title

    def test_author_pulled_from_source(self):
        from docs.paper_2.build_v1_1 import parse_source
        parsed = parse_source(SOURCE_MD)
        assert "Gmys-Casiano" in parsed.author

    def test_all_required_section_ids_present(self):
        from docs.paper_2.build_v1_1 import (
            REQUIRED_SECTION_IDS,
            parse_source,
        )
        parsed = parse_source(SOURCE_MD)
        missing = [sid for sid in REQUIRED_SECTION_IDS
                   if sid not in parsed.sections]
        assert not missing, f"Source markdown missing sections: {missing}"

    def test_subsection_keys_normalized_correctly(self):
        from docs.paper_2.build_v1_1 import parse_source
        parsed = parse_source(SOURCE_MD)
        for sid in ("5.1", "5.2", "6.1", "6.2", "7.1", "7.2", "7.3"):
            assert sid in parsed.sections, f"Missing subsection: {sid}"

    def test_section_one_contains_findings(self):
        """Section 1 must include the Finding 7-11 paragraph."""
        from docs.paper_2.build_v1_1 import parse_source
        parsed = parse_source(SOURCE_MD)
        body = "\n".join(parsed.sections.get("1", []))
        for f in ("Finding 7", "Finding 8", "Finding 9",
                  "Finding 10", "Finding 11"):
            assert f in body, f"Section 1 missing reference to {f}"

    def test_fixture_skeleton_subset_parses(self, tmp_path):
        """A small fixture markdown must parse into the right sections."""
        from docs.paper_2.build_v1_1 import parse_source
        fixture = tmp_path / "fixture.md"
        fixture.write_text(
            "# Title Here\n"
            "\n"
            "**Author Name**\n"
            "Affiliation\n"
            "Date\n"
            "\n"
            "---\n"
            "\n"
            "## Abstract\n"
            "\n"
            "Abstract body.\n"
            "\n"
            "## 1. Intro\n"
            "\n"
            "Intro body.\n"
            "\n"
            "### 1.1 Sub\n"
            "\n"
            "Sub body.\n",
            encoding="utf-8",
        )
        parsed = parse_source(fixture)
        assert parsed.title == "Title Here"
        assert "Author Name" in parsed.author
        assert "Abstract body." in (parsed.sections["abstract"][0])
        assert "Intro body." in (parsed.sections["1"][0])
        assert "Sub body." in (parsed.sections["1.1"][0])


# ----------------------------------------------------------------------------
# Inline rendering
# ----------------------------------------------------------------------------


class TestInlineRendering:
    def test_bold_run_is_bold(self):
        from docx import Document
        from docs.paper_2.build_v1_1 import _render_inline
        doc = Document()
        p = doc.add_paragraph()
        _render_inline(p, "Plain **bold** plain")
        bolds = [r for r in p.runs if r.bold]
        assert len(bolds) == 1
        assert bolds[0].text == "bold"

    def test_italic_run_is_italic(self):
        from docx import Document
        from docs.paper_2.build_v1_1 import _render_inline
        doc = Document()
        p = doc.add_paragraph()
        _render_inline(p, "Plain *em* plain")
        italics = [r for r in p.runs if r.italic]
        assert len(italics) == 1
        assert italics[0].text == "em"

    def test_inline_code_uses_consolas(self):
        from docx import Document
        from docs.paper_2.build_v1_1 import _render_inline
        doc = Document()
        p = doc.add_paragraph()
        _render_inline(p, "See `field_name` here")
        code_runs = [r for r in p.runs if r.font.name == "Consolas"]
        assert len(code_runs) == 1
        assert code_runs[0].text == "field_name"


# ----------------------------------------------------------------------------
# End-to-end build (real source + figures)
# ----------------------------------------------------------------------------


@pytest.fixture(scope="module")
def built_docx(tmp_path_factory) -> Path:
    """Build the v1.1 docx ONCE per test module."""
    from docs.paper_2.build_v1_1 import build_document, parse_source
    parsed = parse_source(SOURCE_MD)
    doc = build_document(parsed, FIGURES_DIR)
    out_dir = tmp_path_factory.mktemp("paper2_v1_1")
    out_path = out_dir / "PAPER2_DRAFT_v1_1.docx"
    doc.save(str(out_path))
    return out_path


class TestEndToEndBuild:
    def test_docx_file_created(self, built_docx):
        assert built_docx.exists()

    def test_docx_size_substantial(self, built_docx):
        # >200KB indicates real prose + 5 inlined PNG figures.
        assert built_docx.stat().st_size > 200 * 1024

    def test_docx_loads_back(self, built_docx):
        from docx import Document
        doc = Document(str(built_docx))
        assert len(doc.paragraphs) > 50

    def test_top_level_headers_present(self, built_docx):
        from docx import Document
        from docs.paper_2.build_skeleton import EXPECTED_TOP_LEVEL_HEADERS
        doc = Document(str(built_docx))
        headings = {p.text for p in doc.paragraphs
                    if p.style.name.startswith("Heading")}
        for h in EXPECTED_TOP_LEVEL_HEADERS:
            assert h in headings, f"Missing top-level header: {h}"

    def test_subsection_headers_present(self, built_docx):
        from docx import Document
        from docs.paper_2.build_skeleton import EXPECTED_SUBSECTION_HEADERS
        doc = Document(str(built_docx))
        headings = {p.text for p in doc.paragraphs
                    if p.style.name.startswith("Heading")}
        for h in EXPECTED_SUBSECTION_HEADERS:
            assert h in headings, f"Missing subsection header: {h}"

    def test_no_todo_markers_remain(self, built_docx):
        from docs.paper_2.number_check import extract_docx_text
        text = extract_docx_text(built_docx)
        assert "[TODO" not in text, "Skeleton TODO marker leaked into v1.1"

    def test_no_missing_prose_markers(self, built_docx):
        from docs.paper_2.number_check import extract_docx_text
        text = extract_docx_text(built_docx)
        assert "[MISSING SOURCE PROSE" not in text, (
            "Some sections did not pull prose from source"
        )

    def test_title_pulled_from_source(self, built_docx):
        from docx import Document
        doc = Document(str(built_docx))
        first_para = doc.paragraphs[0].text
        assert "Deterministic Skeptic" in first_para

    def test_five_figures_inlined(self, built_docx):
        """Every paragraph that contains an inline image counts; should be 5."""
        from docx import Document
        doc = Document(str(built_docx))
        # Inline shapes are accessible from the document directly.
        assert len(doc.inline_shapes) == 5

    def test_references_heading_present(self, built_docx):
        from docx import Document
        doc = Document(str(built_docx))
        headings = [p.text for p in doc.paragraphs
                    if p.style.name.startswith("Heading")]
        assert "References" in headings


# ----------------------------------------------------------------------------
# CLI entry point
# ----------------------------------------------------------------------------


class TestCLI:
    def test_main_returns_zero_with_explicit_paths(self, tmp_path):
        from docs.paper_2.build_v1_1 import main
        out = tmp_path / "v1_1.docx"
        code = main([
            "--source", str(SOURCE_MD),
            "--out", str(out),
            "--figures-dir", str(FIGURES_DIR),
        ])
        assert code == 0
        assert out.exists()
