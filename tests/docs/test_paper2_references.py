"""Tests for docs/paper_2/build_references.py and references.bib.

Covers:
    * The references.bib file parses cleanly.
    * Each expected entry key is present.
    * Author formatting and entry formatting honor ACM/AISec conventions.
    * In-text citations in the v1.1 prose source resolve to bib entries.
    * compile_references appends only beneath an existing References heading.
"""

from __future__ import annotations

from pathlib import Path

import pytest


REPO_ROOT = Path(__file__).resolve().parents[2]
BIB_PATH = REPO_ROOT / "docs/paper_2/references.bib"
SOURCE_MD = REPO_ROOT / "docs/paper_2/source/PAPER2_DRAFT_v1_1_source.md"


EXPECTED_KEYS = (
    "gmys-casiano-2026",
    "berdoz-rugli-wattenhofer-2026",
    "hossain-2025",
    "lee-2024",
    "owasp-2025",
    "sabet-2025",
)


# ----------------------------------------------------------------------------
# Bib parsing
# ----------------------------------------------------------------------------


class TestBibParser:
    def test_real_bib_parses(self):
        from docs.paper_2.build_references import parse_bib_file
        entries = parse_bib_file(BIB_PATH)
        assert len(entries) == 6

    def test_all_expected_keys_present(self):
        from docs.paper_2.build_references import parse_bib_file
        keys = {e.key for e in parse_bib_file(BIB_PATH)}
        for key in EXPECTED_KEYS:
            assert key in keys, f"Missing bib key: {key}"

    def test_every_entry_has_author_and_year(self):
        from docs.paper_2.build_references import parse_bib_file
        for entry in parse_bib_file(BIB_PATH):
            assert entry.get("author"), f"{entry.key} missing author"
            assert entry.get("year"), f"{entry.key} missing year"

    def test_handles_nested_braces_in_author(self):
        """Berdoz et al. has LaTeX accent escapes; parser must tolerate."""
        from docs.paper_2.build_references import parse_bib_file
        entries = {e.key: e for e in parse_bib_file(BIB_PATH)}
        berdoz = entries["berdoz-rugli-wattenhofer-2026"]
        assert "Wattenhofer" in berdoz.get("author")

    def test_inline_minimal_bib_string_parses(self):
        from docs.paper_2.build_references import parse_bib
        text = (
            "@article{foo-2024,\n"
            "  author = {Doe, J. and Roe, R.},\n"
            "  title  = {A Sample Paper},\n"
            "  year   = {2024},\n"
            "  journal = {Sample Journal},\n"
            "}\n"
        )
        entries = parse_bib(text)
        assert len(entries) == 1
        assert entries[0].key == "foo-2024"
        assert entries[0].get("title") == "A Sample Paper"


# ----------------------------------------------------------------------------
# Formatting
# ----------------------------------------------------------------------------


class TestFormatting:
    def test_format_authors_two(self):
        from docs.paper_2.build_references import _format_authors
        assert _format_authors("Doe, J. and Roe, R.") == "Doe, J. and Roe, R."

    def test_format_authors_three(self):
        from docs.paper_2.build_references import _format_authors
        out = _format_authors("Doe, J. and Roe, R. and Smith, S.")
        assert out == "Doe, J., Roe, R., and Smith, S."

    def test_format_authors_strips_latex_accents(self):
        from docs.paper_2.build_references import _format_authors
        assert "Frederic" in _format_authors(r"Berdoz, Fr{\'e}d{\'e}ric")

    def test_format_entry_includes_authors_and_title(self):
        from docs.paper_2.build_references import (
            _strip_latex,
            format_entry,
            parse_bib_file,
        )
        for entry in parse_bib_file(BIB_PATH):
            out = format_entry(entry)
            assert entry.get("year") in out
            # Title may be PLACEHOLDER but the rendered (latex-stripped)
            # leading clause must survive into the output.
            stripped_title = _strip_latex(entry.get("title"))
            assert stripped_title.split(".")[0] in out


# ----------------------------------------------------------------------------
# Citation extraction + mapping
# ----------------------------------------------------------------------------


class TestCitations:
    def test_extract_handles_simple_citation(self):
        from docs.paper_2.build_references import extract_citations
        cites = extract_citations("Foo (Smith, 2024) bar.")
        assert ("Smith", "2024") in cites

    def test_extract_handles_et_al(self):
        from docs.paper_2.build_references import extract_citations
        cites = extract_citations("(Hossain et al., 2025) frobnicates.")
        assert ("Hossain et al.", "2025") in cites

    def test_extract_dedup(self):
        from docs.paper_2.build_references import extract_citations
        text = "(Smith, 2024) ... (Smith, 2024) ..."
        cites = extract_citations(text)
        assert len(cites) == 1

    def test_citation_to_bibkey_simple(self):
        from docs.paper_2.build_references import citation_to_bibkey
        assert citation_to_bibkey(("Smith", "2024")) == "smith-2024"

    def test_citation_to_bibkey_et_al(self):
        from docs.paper_2.build_references import citation_to_bibkey
        assert citation_to_bibkey(("Hossain et al.", "2025")) == "hossain-2025"

    def test_citation_to_bibkey_three_authors(self):
        from docs.paper_2.build_references import citation_to_bibkey
        out = citation_to_bibkey(
            ("Berdoz, Rugli, and Wattenhofer", "2026"),
        )
        assert out == "berdoz-rugli-wattenhofer-2026"

    def test_every_in_text_citation_resolves_to_bib_entry(self):
        """Each citation in the v1.1 prose source must have a bib entry."""
        from docs.paper_2.build_references import (
            citation_to_bibkey,
            extract_citations,
            parse_bib_file,
        )
        prose = SOURCE_MD.read_text(encoding="utf-8")
        bib_keys = {e.key for e in parse_bib_file(BIB_PATH)}
        cites = extract_citations(prose)
        unresolved = [
            (c, citation_to_bibkey(c))
            for c in cites
            if citation_to_bibkey(c) not in bib_keys
        ]
        assert not unresolved, f"Unresolved citations: {unresolved}"


# ----------------------------------------------------------------------------
# Docx integration
# ----------------------------------------------------------------------------


def _docx_with_references_heading(path: Path) -> Path:
    from docx import Document
    doc = Document()
    doc.add_heading("References", level=1)
    doc.save(str(path))
    return path


class TestCompileReferences:
    def test_compile_writes_entries_under_heading(self, tmp_path):
        from docs.paper_2.build_references import compile_references
        target = _docx_with_references_heading(tmp_path / "draft.docx")
        compile_references(target, BIB_PATH)
        from docx import Document
        doc = Document(str(target))
        # After compile, References heading must still be present, plus
        # at least one paragraph below containing a known entry key marker.
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "References" in text
        for key in EXPECTED_KEYS:
            assert key in text, f"Compiled doc missing entry marker [{key}]"

    def test_compile_raises_without_references_heading(self, tmp_path):
        from docx import Document
        from docs.paper_2.build_references import compile_references
        bad = tmp_path / "no-refs.docx"
        Document().save(str(bad))
        with pytest.raises(LookupError):
            compile_references(bad, BIB_PATH)

    def test_compile_raises_when_bib_missing(self, tmp_path):
        from docs.paper_2.build_references import compile_references
        target = _docx_with_references_heading(tmp_path / "draft.docx")
        with pytest.raises(FileNotFoundError):
            compile_references(target, tmp_path / "nope.bib")

    def test_main_returns_zero(self, tmp_path):
        from docs.paper_2.build_references import main
        target = _docx_with_references_heading(tmp_path / "v1_1.docx")
        code = main([
            "--bib", str(BIB_PATH),
            "--docx", str(target),
        ])
        assert code == 0
